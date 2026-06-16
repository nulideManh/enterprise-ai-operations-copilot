import pytest
from app.models.ticket import Ticket
from app.models.email import Email
from app.models.invoice import Invoice
from app.models.github_issue import GitHubIssue
from sqlalchemy import select

@pytest.mark.asyncio
async def test_ticket_agent(async_client, db_session):
    payload = {"issue": "My VPN is down and I cannot connect to the server."}
    response = await async_client.post("/api/agents/ticket", json=payload)
    
    assert response.status_code == 200
    data = response.json()
    assert data["id"] is not None
    assert data["category"] == "Network"
    assert data["priority"] == "High"
    assert "VPN" in data["ticket_summary"]
    
    # Verify DB record creation
    stmt = select(Ticket).where(Ticket.id == data["id"])
    ticket = (await db_session.execute(stmt)).scalar_one_or_none()
    assert ticket is not None
    assert ticket.category == "Network"
    assert ticket.priority == "High"

@pytest.mark.asyncio
async def test_email_agent(async_client, db_session):
    payload = {
        "sender": "partner@company.com",
        "content": "Please review the payment receipt and process the invoice."
    }
    response = await async_client.post("/api/agents/email", json=payload)
    
    assert response.status_code == 200
    data = response.json()
    assert data["id"] is not None
    assert data["category"] == "Finance"
    assert data["confidence"] > 0.7
    
    # Verify DB record creation
    stmt = select(Email).where(Email.id == data["id"])
    email_rec = (await db_session.execute(stmt)).scalar_one_or_none()
    assert email_rec is not None
    assert email_rec.sender == "partner@company.com"
    assert email_rec.category == "Finance"

@pytest.mark.asyncio
async def test_invoice_agent(async_client, db_session):
    # Setup a mock docx/pdf file content simulating an invoice text
    # In invoice_agent, parse_document is called, which handles suffix of filename.
    # We can upload a mock .docx file containing invoice details.
    # Let's create actual docx bytes
    from docx import Document
    from io import BytesIO
    
    doc = Document()
    doc.add_paragraph("Acme Corporation Ltd.")
    doc.add_paragraph("Invoice number: INV-2026-99")
    doc.add_paragraph("Amount: $1,500.00 USD")
    doc.add_paragraph("Date: 2026-06-15")
    
    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    
    files = {"file": ("invoice.docx", file_stream, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
    
    response = await async_client.post("/api/agents/invoice", files=files)
    assert response.status_code == 200
    data = response.json()
    assert data["id"] is not None
    assert data["vendor"] == "Acme Corporation Ltd."
    assert "INV-2026-99" in data["invoice_number"]
    assert "1,500.00" in data["amount"]
    assert data["currency"] == "USD"
    
    # Verify DB record creation
    stmt = select(Invoice).where(Invoice.id == data["id"])
    invoice_rec = (await db_session.execute(stmt)).scalar_one_or_none()
    assert invoice_rec is not None
    assert invoice_rec.vendor == "Acme Corporation Ltd."
    assert "INV-2026-99" in invoice_rec.invoice_number
    assert invoice_rec.amount == data["amount"]

@pytest.mark.asyncio
async def test_github_agent(async_client, db_session):
    payload = {"issue_description": "Fix bug in password validation where spaces are allowed."}
    response = await async_client.post("/api/agents/github", json=payload)
    
    assert response.status_code == 200
    data = response.json()
    assert data["id"] is not None
    assert "validation" in data["root_cause"].lower() or "validation" in data["suggested_fix"].lower()
    assert data["pr_draft"] is not None
    
    # Verify DB record creation
    stmt = select(GitHubIssue).where(GitHubIssue.id == data["id"])
    gh_issue = (await db_session.execute(stmt)).scalar_one_or_none()
    assert gh_issue is not None
    assert "spaces" in gh_issue.issue_description
